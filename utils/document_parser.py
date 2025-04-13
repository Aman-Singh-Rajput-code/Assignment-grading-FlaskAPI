import PyPDF2
import docx
import re
import logging
import traceback

logger = logging.getLogger(__name__)

def extract_text_from_pdf(pdf_path):
    """Extract text content from a PDF file."""
    text = ""
    try:
        logger.info(f"Extracting text from PDF: {pdf_path}")
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            logger.info(f"PDF has {len(pdf_reader.pages)} pages")
            for page_num,page in enumerate(len(pdf_reader.pages)):
#                page = pdf_reader.pages[page_num]
#                page_text = page.extract_text()
#                text += page_text
                page_text = page.extract_text()
                if page_text:
                    text += page_text
#                logger.debug(f"Extracted {len(page_text)} characters from page {page_num+1}")
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        logger.error(traceback.format_exc())
    
#    logger.info(f"Total extracted text length: {len(text)} characters")
#   logger.debug(f"First 200 characters of extracted text: {text[:200]}")
    return text

def extract_text_from_docx(docx_path):
    """Extract text content from a DOCX file."""
    text = ""
    try:
        logger.info(f"Extracting text from DOCX: {docx_path}")
        doc = docx.Document(docx_path)
        logger.info(f"DOCX has {len(doc.paragraphs)} paragraphs")
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        logger.error(traceback.format_exc())
    
    logger.info(f"Total extracted text length: {len(text)} characters")
    logger.debug(f"First 200 characters of extracted text: {text[:200]}")
    return text

def extract_qa_pairs(text):
    """
    Extract question-answer pairs from document text.
    
    The function looks for patterns like:
    Q1: What is the capital of France?
    A1: Paris
    
    Or:
    
    1. What is the capital of France?
    Answer: Paris
    
    Returns a list of dictionaries with question and answer.
    """
    logger.info("Extracting QA pairs from document text")
    qa_pairs = []
    
    # Print the first 500 characters for debugging
    logger.debug(f"First 500 characters of text: {text[:500]}")
    
    # Pattern 1: Q1: ... A1: ...
    pattern1 = r'(?:Q|Question)\s*(\d+)[\s:\.]*(.*?)(?:\n|\r\n?)(?:A|Answer)\s*\1[\s:\.]*(.*?)(?=(?:Q|Question)\s*\d+|\Z)'
    logger.debug(f"Trying pattern 1: {pattern1}")
    
    # Pattern 2: 1. Question ... Answer: ...
    pattern2 = r'(\d+)[\.:\)]\s*(.*?)(?:\n|\r\n?)(?:A|Answer|Ans)[\.:\s]*(.*?)(?=\d+[\.:\)]\s*|\Z)'
    logger.debug(f"Trying pattern 2: {pattern2}")
    
    # Try the first pattern
    matches = re.findall(pattern1, text, re.DOTALL | re.IGNORECASE)
    logger.debug(f"Pattern 1 found {len(matches)} matches")
    
    if matches:
        for match in matches:
            qa_pair = {
                'question_num': match[0],
                'question': match[1].strip(),
                'answer': match[2].strip()
            }
            qa_pairs.append(qa_pair)
            logger.debug(f"Found QA pair: Q{match[0]}: {match[1][:30]}... A: {match[2][:30]}...")
    else:
        # Try the second pattern
        matches = re.findall(pattern2, text, re.DOTALL | re.IGNORECASE)
        logger.debug(f"Pattern 2 found {len(matches)} matches")
        
        if matches:
            for match in matches:
                qa_pair = {
                    'question_num': match[0],
                    'question': match[1].strip(),
                    'answer': match[2].strip()
                }
                qa_pairs.append(qa_pair)
                logger.debug(f"Found QA pair: Q{match[0]}: {match[1][:30]}... A: {match[2][:30]}...")
    
    # If no structured format is found, try a more general approach
    if not qa_pairs:
        logger.info("No structured QA pairs found with regex, trying line-by-line approach")
        lines = text.split('\n')
        current_question = None
        current_q_num = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # Print some sample lines for debugging
            if i < 10:
                logger.debug(f"Line {i}: {line}")
            
            # Look for question indicators
            q_match = re.match(r'(?:Q|Question)\s*(\d+)[\.:\s]*(.+)', line, re.IGNORECASE)
            if q_match:
                current_q_num = q_match.group(1)
                current_question = q_match.group(2).strip()
                logger.debug(f"Found question {current_q_num}: {current_question[:30]}...")
                continue
                
            # Another question pattern
            q_match = re.match(r'(\d+)[\.:\)]\s*(.+)', line, re.IGNORECASE)
            if q_match and not current_question:
                current_q_num = q_match.group(1)
                current_question = q_match.group(2).strip()
                logger.debug(f"Found question {current_q_num}: {current_question[:30]}...")
                continue
            
            # Look for answer indicators
            if current_question and re.match(r'(?:A|Answer|Ans)[\.:\s]*(.+)', line, re.IGNORECASE):
                answer = re.match(r'(?:A|Answer|Ans)[\.:\s]*(.+)', line, re.IGNORECASE).group(1).strip()
                qa_pair = {
                    'question_num': current_q_num,
                    'question': current_question,
                    'answer': answer
                }
                qa_pairs.append(qa_pair)
                logger.debug(f"Found answer for Q{current_q_num}: {answer[:30]}...")
                current_question = None
                current_q_num = None
    
    logger.info(f"Extracted {len(qa_pairs)} QA pairs in total")
    return qa_pairs